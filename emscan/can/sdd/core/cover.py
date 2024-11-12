try:
    from ....config import PATH
    from ....svn.vcon import VersionControl
    from ...db.db import DB
except ImportError:
    from emscan.config import PATH
    from emscan.svn.vcon import VersionControl
    from emscan.can.db.db import DB
from datetime import datetime
from docx import Document
from docx.parts.styles import Styles
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt, Inches
import os

class Cover:
    __version:str = ""
    BASEDB   :str = "자체제어기_KEFICO-EMS_CANFD.xlsx"
    DIVISION :str = "차량제어솔루션팀 / Vehicle Control Solution Team"
    def __init__(self, doc:Document):
        self.doc:Document = doc
        return

    @property
    def titleStyle(self) -> Styles:
        if not hasattr(self, "_title"):
            style = self.doc.styles.add_style("title", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 Head"
            style.font.size = Pt(28)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.__setattr__("_title", style)
        return self.__getattribute__("_title")

    @property
    def overviewLeftStyle(self) -> Styles:
        if not hasattr(self, "_ovl"):
            style = self.doc.styles.add_style("overview_left", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 Text"
            style.font.size = Pt(11)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.__setattr__("_ovl", style)
        return self.__getattribute__("_ovl")

    @property
    def overviewRightStyle(self) -> Styles:
        if not hasattr(self, "_ovr"):
            style = self.doc.styles.add_style("overview_right", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 Text"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.__setattr__("_ovr", style)
        return self.__getattribute__("_ovr")

    @property
    def version(self) -> str:
        return self.__version

    @version.setter
    def version(self, version:str):
        self.__version = version

    def setTitle(self):
        self.doc.add_paragraph("\nEMS/ASW CAN-FD SPECIFICATION\n\n", style=self.titleStyle)
        return

    def setCI(self):
        ci = os.path.join(os.path.dirname(os.path.dirname(__file__)), r"archive/ci_cover.png")
        paragraph = self.doc.add_paragraph()
        runner = paragraph.add_run()
        runner.add_picture(ci, width=Inches(6))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.doc.add_paragraph("\n\n", style=self.titleStyle)
        return


    def setOverview(self):
        log = VersionControl.log(PATH.SVN.CAN.DB.file(self.BASEDB))
        self.version = log.iloc[0, 0]
        items = {
            'DATABASE': DB.traceability.split("_V")[0],
            'REVISION': self.version,
            'DIVISION': self.DIVISION,
            'RELEASE' : datetime.now().strftime("%Y-%m-%d"),
            'HISTORY' : '\n'.join(log["revision"] + ' @' + log["datetime"] + ' ' + log["log"])
        }
        vWidth = self.doc.sections[0].page_width
                 # - self.doc.sections[0].left_margin \
                 # - self.doc.sections[0].right_margin \
                 # - Inches(1)
        table = self.doc.add_table(rows=len(items), cols=2)
        table.style = 'Table Grid'
        for n, (key, value) in enumerate(items.items()):
            left = table.rows[n].cells[0]
            left.width = Inches(1)
            name = left.paragraphs[0]
            name.style = self.overviewLeftStyle
            name.text = key

            right = table.rows[n].cells[1]
            right.width = vWidth
            text = right.paragraphs[0]
            text.style = self.overviewRightStyle
            text.text = value
        return