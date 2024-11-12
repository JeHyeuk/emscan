from docx import Document
from docx.parts.styles import Styles
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, ns
from datetime import datetime



class Section:
    def __init__(self, doc:Document):
        self.doc:Document = doc
        return

    @property
    def footerStyle(self) -> Styles:
        if not hasattr(self, "_footer_style"):
            style = self.doc.styles.add_style("_footer_style", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(10)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.__setattr__("_footer_style", style)
        return self.__getattribute__("_footer_style")

    @property
    def headerStyleLeft(self) -> Styles:
        if not hasattr(self, "_header_style_left"):
            style = self.doc.styles.add_style("_header_style_left", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(8)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            self.__setattr__("_header_style_left", style)
        return self.__getattribute__("_header_style_left")

    @property
    def headerStyleRight(self) -> Styles:
        if not hasattr(self, "_header_style_right"):
            style = self.doc.styles.add_style("_header_style_right", WD_PARAGRAPH_ALIGNMENT.CENTER)
            style.font.name = "현대산스 text"
            style.font.size = Pt(8)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = False
            style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            self.__setattr__("_header_style_right", style)
        return self.__getattribute__("_header_style_right")

    def setMargin(self):
        section = self.doc.sections[0]
        section.left_margin = \
        section.right_margin = \
        section.bottom_margin = Inches(0.5)
        section.top_margin = Inches(1)
        return

    def setHeader(self, version:str):
        header = self.doc.sections[0].header
        for paragraph in header.paragraphs:
            p = getattr(paragraph, "_element")
            p.getparent().remove(p)

        table = header.add_table(rows=1, cols=3, width=self.doc.sections[0].page_width)

        left = table.rows[0].cells[0].paragraphs[0]
        left.text = f"EMS CAN/CANFD\nDOC {version}"
        left.style = self.headerStyleLeft

        right = table.rows[0].cells[2].paragraphs[0]
        right.text = f"Vehicle Control Solution\nHYUNDAI KEFICO Co.,Ltd."
        right.style = self.headerStyleRight
        return

    def setFooter(self):
        footer = self.doc.sections[0].footer.paragraphs[0]
        footer.text = f"Copyright(c) 2020-{datetime.now().year} " \
                      f"HYUNDAI-KEFICO Co., Ltd. All Rights Reserved"
        footer.style = self.footerStyle
        return
